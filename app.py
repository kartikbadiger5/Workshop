import streamlit as st
import PyPDF2
import docx
import google.generativeai as genai
from fpdf import FPDF
import pymongo
import hashlib
from pymongo import MongoClient

# Configure page layout
st.set_page_config(layout="wide", page_title="AI Resume Optimizer")

# Configure Gemini
genai.configure(api_key="AIzaSyAkUOKHDBVJS0zx577VYHLPdMnUDj2E6Jk")
model = genai.GenerativeModel('gemini-1.5-flash')

# MongoDB Atlas connection
def get_mongo_client():
    # Replace with your MongoDB Atlas connection string
    mongo_uri = "mongodb+srv://kartik:2002@cluster0.2bok0.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"
    client = MongoClient(mongo_uri)
    return client

# Initialize MongoDB collection
def init_db():
    client = get_mongo_client()
    db = client["resume_optimizer"]
    users_collection = db["users"]
    return users_collection

# Create a new user
def create_user(username, password):
    users_collection = init_db()
    hashed_password = hashlib.sha256(password.encode()).hexdigest()
    user = {"username": username, "password": hashed_password, "first_time": True}
    try:
        users_collection.insert_one(user)
        return True
    except pymongo.errors.DuplicateKeyError:
        return False

# Verify user credentials
def verify_user(username, password):
    users_collection = init_db()
    hashed_password = hashlib.sha256(password.encode()).hexdigest()
    user = users_collection.find_one({"username": username, "password": hashed_password})
    return user is not None

# Check if it's the user's first time
def is_first_time(username):
    users_collection = init_db()
    user = users_collection.find_one({"username": username})
    if user and user.get("first_time", True):
        return True
    return False

# Update user's first_time flag to False
def update_first_time(username):
    users_collection = init_db()
    users_collection.update_one({"username": username}, {"$set": {"first_time": False}})

# Session state for user authentication
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'username' not in st.session_state:
    st.session_state['username'] = None
if 'first_time' not in st.session_state:
    st.session_state['first_time'] = True

def extract_text(file):
    """Extract text from PDF or DOCX files"""
    if file.type == "application/pdf":
        reader = PyPDF2.PdfReader(file)
        return "\n".join(page.extract_text() for page in reader.pages)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        return "\n".join(para.text for para in doc.paragraphs)
    return file.getvalue().decode("utf-8")

def analyze_resume(resume_text, job_desc_text):
    """Use Gemini to analyze skill gaps and suggest improvements with proper formatting"""
    prompt = f"""
    Analyze this resume against the job description and provide:
    1. Skill gaps
    2. Suggested improvements
    3. Optimized resume text in proper resume format
    
    The optimized resume must include these sections with proper formatting:
    - Name and Contact Information
    - Professional Summary
    - Skills (aligned with job requirements)
    - Work Experience (with relevant achievements)
    - Education
    - Certifications (if any)
    
    Maintain proper formatting with:
    - Clear section headings
    - Bullet points for achievements
    - Consistent spacing
    - Professional language
    
    Resume:
    {resume_text}
    
    Job Description:
    {job_desc_text}
    """
    response = model.generate_content(prompt)
    return response.text

def create_pdf(content):
    """Create a PDF file from text with built-in Unicode support"""
    pdf = FPDF()
    pdf.add_page()
    # Use built-in font with Unicode support
    pdf.set_font('helvetica', '', 12)
    # Add content with proper encoding
    pdf.multi_cell(0, 10, txt=content.encode('latin1', 'replace').decode('latin1'))
    return pdf.output(dest='S').encode('latin1')

def create_docx(content):
    """Create a DOCX file from text"""
    doc = docx.Document()
    doc.add_paragraph(content)
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def login_form():
    with st.form("Login"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submit_button = st.form_submit_button("Login")
        if submit_button:
            if verify_user(username, password):
                st.session_state['logged_in'] = True
                st.session_state['username'] = username
                st.session_state['first_time'] = is_first_time(username)
                st.success("Logged in successfully!")
            else:
                st.error("Invalid username or password")

def register_form():
    with st.form("Register"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submit_button = st.form_submit_button("Register")
        if submit_button:
            if create_user(username, password):
                st.success("Registration successful! Please login.")
            else:
                st.error("Username already exists")

def main():
    # Create a two-column layout
    left_col, right_col = st.columns([1, 2])
    
    with left_col:
        st.image("https://img.icons8.com/color/96/000000/resume.png", width=80)
        st.title("Resume Optimizer")
        st.markdown("---")
        
        # Check if it's the user's first time or if they are logged in
        if not st.session_state['logged_in'] and not st.session_state['first_time']:
            st.subheader("Login Required")
            login_form()
            st.markdown("---")
            register_form()
            return
        
        # File upload section
        st.subheader("📄 Upload Files")
        
        # File uploaders with custom styling
        resume_file = st.file_uploader(
            "Upload Resume", 
            type=["pdf", "docx"],
            help="Supported formats: PDF, DOCX"
        )
        
        job_desc_file = st.file_uploader(
            "Upload Job Description", 
            type=["pdf", "docx", "txt"],
            help="Supported formats: PDF, DOCX, TXT"
        )
        
        # Status indicators
        if resume_file:
            st.success(f"✅ Resume uploaded: {resume_file.name}")
        else:
            st.info("⏳ Waiting for resume...")
            
        if job_desc_file:
            st.success(f"✅ Job description uploaded: {job_desc_file.name}")
        else:
            st.info("⏳ Waiting for job description...")
        
        # Analyze button (only shown when both files are uploaded)
        if resume_file and job_desc_file:
            st.markdown("---")
            analyze_button = st.button("🔍 Analyze Resume", type="primary", use_container_width=True)
        else:
            analyze_button = False
            
    with right_col:
        if not (resume_file and job_desc_file):
            # Display welcome message when files are not uploaded
            st.markdown("""
            ## Welcome to AI Resume Optimizer! 👋
            
            This tool helps you optimize your resume for specific job positions using AI.
            
            ### How it works:
            1. Upload your resume (PDF/DOCX)
            2. Upload the job description (PDF/DOCX/TXT)
            3. Get an AI analysis of skill gaps and suggestions
            4. Download your optimized resume
            
            ### Upload your files on the left panel to get started!
            """)
            
            # Display sample resume structure
            with st.expander("View Sample Resume Structure"):
                st.markdown("""
                ### John Doe
                john.doe@email.com | (123) 456-7890 | linkedin.com/in/johndoe
                
                ### Professional Summary
                Experienced software engineer with 5+ years specializing in web development...
                
                ### Skills
                - Programming: JavaScript, Python, Java
                - Frameworks: React, Node.js, Django
                - Tools: Git, Docker, AWS
                
                ### Work Experience
                **Senior Software Engineer** | ABC Company | Jan 2020 - Present
                - Developed and maintained enterprise web applications
                - Led team of 5 developers on critical projects
                
                ### Education
                **Bachelor of Science in Computer Science** | University Name | 2015-2019
                
                ### Certifications
                - AWS Certified Developer
                - Scrum Master Certification
                """)
        
        elif analyze_button:
            # Show a spinner while processing
            with st.spinner("AI is analyzing your resume against the job description..."):
                # Extract text
                resume_text = extract_text(resume_file)
                job_desc_text = extract_text(job_desc_file)
                
                # Analyze with Gemini
                analysis = analyze_resume(resume_text, job_desc_text)
            
            # Display results in tabs
            tab1, tab2 = st.tabs(["Analysis Results", "Download Options"])
            
            with tab1:
                st.markdown("## AI Analysis Results")
                st.markdown(analysis)
            
            with tab2:
                st.markdown("## Download Your Optimized Resume")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.download_button(
                        label="📄 Download as TXT",
                        data=analysis.encode('utf-8'),
                        file_name="optimized_resume.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                with col2:
                    st.download_button(
                        label="📑 Download as PDF",
                        data=create_pdf(analysis),
                        file_name="optimized_resume.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                with col3:
                    st.download_button(
                        label="📝 Download as DOCX",
                        data=create_docx(analysis),
                        file_name="optimized_resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            # Update first_time flag after the first use
            if st.session_state['first_time']:
                update_first_time(st.session_state['username'])
                st.session_state['first_time'] = False

if __name__ == "__main__":
    main()